import asyncio
import os
import sys
from pathlib import Path
from typing import Optional, List, Dict, Any
import json
from notebooklm import NotebookLMClient, Notebook
from notebooklm.auth import AuthTokens
from pptx import Presentation
import pdfkit
from notebooklm.rpc.types import ReportFormat, QuizDifficulty, QuizQuantity
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Wrapper to adapt Playwright APIResponse to httpx.Response interface
class PlaywrightResponseAdapter:
    def __init__(self, pw_response, body_text):
        self._pw_response = pw_response
        self._text = body_text
        
    @property
    def status_code(self):
        return self._pw_response.status
        
    @property
    def text(self):
        return self._text
        
    @property
    def url(self):
        return self._pw_response.url
        
    @property
    def headers(self):
        return self._pw_response.headers
        
    def raise_for_status(self):
        if 400 <= self.status_code < 600:
            raise Exception(f"HTTP Error {self.status_code}: {self.text[:200]}")
            
    def json(self):
        return json.loads(self._text)

# Wrapper to adapt Playwright APIRequestContext to httpx.AsyncClient interface
class PlaywrightHttpClient:
    def __init__(self, request_context):
        self.request = request_context
        self.headers = {} # Mock headers storage that library might try to update

    async def post(self, url, content=None, headers=None, **kwargs):
        # Merge mocked headers with request-specific headers
        # We start with required headers for Google RPC
        req_headers = {
            "Content-Type": "application/x-www-form-urlencoded;charset=UTF-8",
            "Origin": "https://notebooklm.google.com",
            "Referer": "https://notebooklm.google.com/",
            "X-Same-Domain": "1",
        }
        
        # Merge in any library-provided headers (like Cookie, if we wanted them)
        merged_headers = {**self.headers, **(headers or {})}
        
        # CRITICAL: Remove 'Cookie' from headers and let Playwright Context handle it.
        # Manually sending cookies might conflict with the browser's internal jar
        # or be missing secure/http-only cookies that the browser handles.
        if "Cookie" in merged_headers:
            del merged_headers["Cookie"]
            
        final_headers = {**req_headers, **merged_headers}
        
        # Handle timeout: httpx sends seconds (float), Playwright wants ms (int) or 0
        timeout_sec = kwargs.get("timeout", 300.0) # Default to 5 min if not specified
        if timeout_sec is None:
            timeout_ms = 0
        else:
            timeout_ms = int(timeout_sec * 1000)

        print(f"  [DEBUG] PlaywrightHttpClient.post timeout: {timeout_sec}s -> {timeout_ms}ms") # Debug logging

        # Playwright expects 'data' for body
        response = await self.request.post(url, data=content, headers=final_headers, timeout=timeout_ms)
        text = await response.text()
        return PlaywrightResponseAdapter(response, text)

    async def get(self, url, headers=None, **kwargs):
        final_headers = {**self.headers, **(headers or {})}
        
        # Handle timeout
        timeout_sec = kwargs.get("timeout", 300.0)
        if timeout_sec is None:
            timeout_ms = 0
        else:
            timeout_ms = int(timeout_sec * 1000)

        print(f"  [DEBUG] PlaywrightHttpClient.get timeout: {timeout_sec}s -> {timeout_ms}ms") # Debug logging
            
        response = await self.request.get(url, headers=final_headers, timeout=timeout_ms)
        text = await response.text()
        return PlaywrightResponseAdapter(response, text)
        
    async def aclose(self):
        pass # Browser is managed by manager

# Global state manager
class NotebookManager:
    def __init__(self):
        self.client: Optional[NotebookLMClient] = None
        self.auth: Optional[AuthTokens] = None
        self.current_notebook_id: Optional[str] = None
        self.client: Optional[NotebookLMClient] = None
        self.auth: Optional[AuthTokens] = None
        self.current_notebook_id: Optional[str] = None
        self.chat_history: Dict[str, List[Dict]] = {}
        self.artifacts_store: Dict[str, List[Dict]] = {} # Persistent artifacts
        
        # Playwright objects
        self.playwright = None
        self.browser = None
        self.context = None
        self.page = None
        
        self._load_history()

    def _load_history(self):
        history_file = Path("chat_history.json")
        if history_file.exists():
            try:
                data = json.loads(history_file.read_text(encoding="utf-8"))
                if isinstance(data, list):
                    # Migrate old list format to dict under 'default' or migrate later
                    self.chat_history = {"legacy": data}
                else:
                    self.chat_history = data
            except:
                self.chat_history = {}

        # Load artifacts if stored in the same file or a separate one
        # For simplicity, we'll assume chat_history.json might contain an "artifacts" key in the future
        # But for now, let's keep it backward compatible or use a separate key if history structure allows
        # Since self.chat_history is a dict of lists, we might need a separate file or a special key
        
        artifacts_file = Path("artifacts.json")
        if artifacts_file.exists():
            try:
                self.artifacts_store = json.loads(artifacts_file.read_text(encoding="utf-8"))
            except:
                self.artifacts_store = {}

    def _save_history(self):
        Path("chat_history.json").write_text(
            json.dumps(self.chat_history, indent=2, ensure_ascii=False), 
            encoding="utf-8"
        )
        Path("artifacts.json").write_text(
            json.dumps(self.artifacts_store, indent=2, ensure_ascii=False),
            encoding="utf-8"
        )

    def add_message(self, notebook_id: str, role: str, text: str):
        if notebook_id not in self.chat_history:
            self.chat_history[notebook_id] = []
        self.chat_history[notebook_id].append({"role": role, "text": text})
        self._save_history()

    def add_artifact(self, notebook_id: str, type: str, title: str, details: Dict):
        if notebook_id not in self.artifacts_store:
            self.artifacts_store[notebook_id] = []
        
        # Add timestamp
        from datetime import datetime
        artifact = {
            "id": f"{type}_{int(datetime.now().timestamp())}",
            "type": type,
            "title": title,
            "details": details, # e.g. filename, instructions, content summary
            "created_at": datetime.now().isoformat()
        }
        self.artifacts_store[notebook_id].append(artifact)
        self._save_history()
        return artifact

    def get_artifacts(self, notebook_id: str):
        return self.artifacts_store.get(notebook_id, [])

    def get_artifact_content(self, notebook_id: str, artifact_id: str):
        artifacts = self.artifacts_store.get(notebook_id, [])
        artifact = next((a for a in artifacts if a["id"] == artifact_id), None)
        if not artifact:
            return None
        
        path = artifact["details"].get("path")
        if not path or not os.path.exists(path):
            return None
        
        # Determine how to read based on type/extension
        try:
            return Path(path).read_text(encoding="utf-8")
        except Exception as e:
            print(f"Error reading artifact {path}: {e}")
            return None

    def get_history(self, notebook_id: str) -> List[Dict]:
        return self.chat_history.get(notebook_id, [])
    async def query(self, prompt: str):
        if not self.client:
            raise Exception("Not authenticated")
        if not self.current_notebook_id:
            raise Exception("No notebook selected")
        
        self.add_message(self.current_notebook_id, "user", prompt)
        result = await self.client.chat.ask(self.current_notebook_id, prompt)
        
        # Handle result content
        answer_text = result.answer
        
        self.add_message(self.current_notebook_id, "ai", answer_text)
        return answer_text

    async def _launch_browser(self, headless=True):
        """Launch browser with persistent context to save login state"""
        from playwright.async_api import async_playwright
        
        # Detect browser executable path
        browser_executable = None
        
        if getattr(sys, 'frozen', False):
            # Running from PyInstaller bundle
            bundle_dir = sys._MEIPASS
            print(f"Running from PyInstaller bundle: {bundle_dir}")
            
            # Try to find bundled browser
            playwright_browsers_path = os.path.join(bundle_dir, 'playwright_browsers')
            if os.path.exists(playwright_browsers_path):
                # Look for chromium directory
                chromium_dirs = [d for d in os.listdir(playwright_browsers_path) if d.startswith('chromium-')]
                if chromium_dirs:
                    browser_executable = os.path.join(
                        playwright_browsers_path, 
                        chromium_dirs[0], 
                        'chrome-win64', 
                        'chrome.exe'
                    )
                    print(f"Found bundled browser at: {browser_executable}")
        
        # If not found in bundle or not frozen, use system Playwright browsers
        if not browser_executable or not os.path.exists(browser_executable):
            # Reset to None to ensure clean search state
            browser_executable = None
            
            print("Searching for system Playwright browsers...")
            # Manual search for system Playwright browsers to avoid 'frozen' bias
            import platform
            import subprocess
            
            def find_system_browser():
                if platform.system() == 'Windows':
                    # 1. Search Playwright managed browsers
                    search_paths = [
                        os.path.join(os.environ.get('LOCALAPPDATA', ''), 'ms-playwright'),
                        os.path.join(os.environ.get('USERPROFILE', ''), 'AppData', 'Local', 'ms-playwright')
                    ]
                    for base_path in search_paths:
                        if os.path.exists(base_path):
                            try:
                                print(f"Checking system path: {base_path}")
                                # Verify if any chromium version exists
                                for d in os.listdir(base_path):
                                    if d.startswith('chromium-'):
                                        candidate = os.path.join(base_path, d, 'chrome-win64', 'chrome.exe')
                                        if os.path.exists(candidate):
                                            print(f"Found system browser manually at: {candidate}")
                                            return candidate
                            except Exception as e:
                                print(f"Error searching path {base_path}: {e}")
                                
                    # 2. Search System Browsers (Chrome/Edge)
                    # Playwright can use standard Chrome/Edge executables
                    common_paths = [
                        os.path.expandvars(r"%ProgramFiles%\Google\Chrome\Application\chrome.exe"),
                        os.path.expandvars(r"%ProgramFiles(x86)%\Google\Chrome\Application\chrome.exe"),
                        os.path.expandvars(r"%ProgramFiles%\Microsoft\Edge\Application\msedge.exe"),
                        os.path.expandvars(r"%ProgramFiles(x86)%\Microsoft\Edge\Application\msedge.exe"),
                    ]
                    for path in common_paths:
                        if os.path.exists(path):
                             print(f"Found system browser (Chrome/Edge) at: {path}")
                             return path
                             
                return None

            # First attempt to find browser
            browser_executable = find_system_browser()
            
            # If not found, attempt to install
            if not browser_executable:
                print("Browser not found. Attempting to install Chromium...")
                try:
                    from playwright._impl._driver import compute_driver_executable
                    driver_exec, driver_env = compute_driver_executable()
                    
                    if os.path.exists(driver_exec):
                        print(f"Installing browser using driver: {driver_exec}")
                        
                        install_cmd = [driver_exec, "install", "chromium"]
                        
                        # Configure startup info to hide window on Windows
                        startupinfo = None
                        if platform.system() == 'Windows':
                            startupinfo = subprocess.STARTUPINFO()
                            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                            
                        subprocess.run(install_cmd, env=driver_env, check=True, startupinfo=startupinfo)
                        print("Browser installation completed successfully.")
                        
                        # Search again after installation
                        browser_executable = find_system_browser()
                    else:
                        print(f"Driver executable not found at: {driver_exec}")
                except Exception as e:
                    print(f"Failed to install browser: {e}")
            
            # Use default Playwright detection as last resort fallback
            if not browser_executable:
                try:
                    # Start playwright temporarily to get executable path
                    temp_pw = await async_playwright().start()
                    browser_executable = temp_pw.chromium.executable_path
                    await temp_pw.stop()
                    print(f"Using system Playwright browser (API) at: {browser_executable}")
                except Exception as e:
                    print(f"Warning: Could not detect browser path via API: {e}")
                    browser_executable = None
        
        if self.playwright: 
            await self.playwright.stop()

        self.playwright = await async_playwright().start()
        
        # Use a local directory to store browser profile (cookies, etc.)
        user_data_dir = os.path.abspath("chrome_profile")
        if not os.path.exists(user_data_dir):
            os.makedirs(user_data_dir)
            
        print(f"Launching browser (Headless: {headless}) with profile at: {user_data_dir}")
        if browser_executable:
            print(f"Using browser executable: {browser_executable}")
        
        try:
            # Prepare launch arguments
            launch_args = {
                "user_data_dir": user_data_dir,
                "headless": headless,
                "args": [
                    "--disable-blink-features=AutomationControlled",
                    "--no-sandbox", 
                    "--disable-infobars"
                ],
                "ignore_default_args": ["--enable-automation"],
                "user_agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "viewport": {"width": 1280, "height": 720}
            }
            
            # Add executable_path if we found one
            if browser_executable and os.path.exists(browser_executable):
                launch_args["executable_path"] = browser_executable
            
            self.context = await self.playwright.chromium.launch_persistent_context(**launch_args)
            
            # Stealth scripts
            await self.context.add_init_script("""
                Object.defineProperty(navigator, 'webdriver', {
                    get: () => undefined
                });
            """)
            
            self.page = self.context.pages[0] if self.context.pages else await self.context.new_page()
            
        except Exception as e:
            print(f"Error launching browser: {e}")
            if self.playwright: await self.playwright.stop()
            raise e

    async def _initialize_client(self):
        """Initialize NotebookLM client with current browser session"""
        print("Extracting tokens from page...")
        content = await self.page.content()
        current_url = self.page.url
        
        cookies = await self.context.cookies()
        cookie_dict = {c['name']: c['value'] for c in cookies}
        
        from notebooklm.auth import extract_csrf_from_html, extract_session_id_from_html
        
        # We might need to handle cases where token extraction fails
        try:
            csrf_token = extract_csrf_from_html(content, current_url)
            session_id = extract_session_id_from_html(content, current_url)
        except Exception as e:
            print(f"Token extraction failed: {e}")
            return False

        print("Initializing NotebookLM Client...")
        self.auth = AuthTokens(cookies=cookie_dict, csrf_token=csrf_token, session_id=session_id)
        self.client = NotebookLMClient(auth=self.auth)
        
        # Inject Playwright HTTP Client
        pw_http_client = PlaywrightHttpClient(self.page.context.request)
        self.client._core._http_client = pw_http_client

        # CRITICAL: Save storage state to the library's expected location
        # This fixes download methods (audio/video/etc.) which look for this file
        try:
            storage_path = Path.home() / ".notebooklm" / "storage_state.json"
            storage_path.parent.mkdir(parents=True, exist_ok=True)
            await self.context.storage_state(path=str(storage_path))
            print(f"Saved storage state to {storage_path}")
        except Exception as e:
            print(f"Warning: Failed to save storage state: {e}")

        return True

    async def login_with_playwright(self):
        """Interactive login with persistent context"""
        # Close existing headless session if any
        if self.context:
            await self.context.close()
            
        await self._launch_browser(headless=False)
        
        print("Navigating to NotebookLM...")
        await self.page.goto("https://notebooklm.google.com/")
        
        print("Waiting for login...")
        # Wait for user to be logged in (url is notebooklm.google.com and NOT accounts.google.com)
        # We'll wait until we can find the home page indicator or URL matches
        
        try:
            # wait up to 5 minutes for login
            await self.page.wait_for_url("https://notebooklm.google.com/", timeout=300000, wait_until="domcontentloaded")
            
            # Wait loop to ensure we are really logged in and not redirected back
            for _ in range(60):
                if "accounts.google.com" not in self.page.url:
                    break
                await asyncio.sleep(1)
            
            await self.page.wait_for_timeout(2000) # Settle parameters
            
            if await self._initialize_client():
                print("Login successful locally. Switching to headless background session...")
                
                # Explicitly close the visible browser to save profile
                await self.context.close()
                await self.playwright.stop()
                self.playwright = None
                
                # Relaunch in headless mode for background operation
                print("Relaunching headless...")
                await self._launch_browser(headless=True)
                
                # Navigate to site to restore session/tokens
                print("Navigating to restore session...")
                try:
                    await self.page.goto("https://notebooklm.google.com/")
                    await self.page.wait_for_url("https://notebooklm.google.com/", timeout=60000, wait_until="domcontentloaded")
                    await self.page.wait_for_timeout(2000) # Wait for hydration
                except Exception as e:
                    print(f"Error restoring session navigation: {e}")
                    return False
                
                # Re-bind client to new headless context
                if await self._initialize_client():
                    print("Headless session initialized successfully.")
                    return True
                else:
                    print("Failed to initialize headless session.")
                    return False
            else:
                return False
                
        except Exception as e:
            print(f"Login failed or timed out: {e}")
            return False

    async def try_auto_connect(self):
        """Try to auto-connect using persistent profile"""
        print("Attempting auto-connect with persistent profile...")
        try:
            await self._launch_browser(headless=True)
            
            print("Navigating to check session...")
            await self.page.goto("https://notebooklm.google.com/")
            await self.page.wait_for_timeout(2000)
            
            if "accounts.google.com" in self.page.url:
                print("Auto-connect failed: Redirected to login page.")
                await self.context.close()
                self.context = None
                return False
                
            print("Session appears valid.")
            if await self._initialize_client():
                print("Auto-connect successful!")
                return True
            else:
                await self.context.close()
                self.context = None
                return False
                
        except Exception as e:
            print(f"Auto-connect error: {e}")
            if self.context:
                await self.context.close()
            self.context = None
            return False

    # Compatibility method (less used now)
    async def login(self, cookies: Dict[str, str] = None, cookie_str: str = None):
        print("Manual login deprecated in favor of Playwright flow")
        return False

    async def list_notebooks(self) -> List[Notebook]:
        if not self.client:
            await self.try_auto_connect()
        if not self.client:
            raise Exception("Not authenticated")
        return await self.client.notebooks.list()

    async def create_notebook(self, title: str) -> Notebook:
        if not self.client: raise Exception("Not authenticated")
        return await self.client.notebooks.create(title)

    async def rename_notebook(self, notebook_id: str, new_title: str):
        if not self.client: raise Exception("Not authenticated")
        # Try finding the correct method on the client
        # Assuming typical REST-like method structure
        if hasattr(self.client.notebooks, 'rename'):
            await self.client.notebooks.rename(notebook_id, new_title)
        elif hasattr(self.client.notebooks, 'update'):
             await self.client.notebooks.update(notebook_id, title=new_title)
        else:
            # Fallback: Raise error if not supported, or let's try direct call
             await self.client.notebooks.rename(notebook_id, new_title)

    async def delete_notebook(self, notebook_id: str):
        if not self.client: raise Exception("Not authenticated")
        # If we act on current notebook, clear selection
        if self.current_notebook_id == notebook_id:
            self.current_notebook_id = None
        await self.client.notebooks.delete(notebook_id)

    async def add_source_url(self, notebook_id: str, url: str):
        if not self.client: raise Exception("Not authenticated")
        return await self.client.sources.add_url(notebook_id, url)

    async def add_source_text(self, notebook_id: str, title: str, content: str):
        if not self.client: raise Exception("Not authenticated")
        return await self.client.sources.add_text(notebook_id, title, content)

    async def add_source_file(self, notebook_id: str, file_path: str):
        if not self.client: raise Exception("Not authenticated")
        return await self.client.sources.add_file(notebook_id, file_path)

    async def delete_source(self, notebook_id: str, source_id: str):
        if not self.client: raise Exception("Not authenticated")
        await self.client.sources.delete(notebook_id, source_id)

    def set_notebook(self, notebook_id: str):
        self.current_notebook_id = notebook_id

    async def get_sources(self) -> List[Dict[str, Any]]:
        if not self.client:
            raise Exception("Not authenticated")
        if not self.current_notebook_id:
            raise Exception("No notebook selected")
            
        sources = await self.client.sources.list(self.current_notebook_id)
        print(f"DEBUG: Found {len(sources)} sources")
        
        # Serialize to dicts
        serialized = []
        for s in sources:
            try:
                s_type = "unknown"
                if hasattr(s, "kind"):
                    # Use .name if available (Enum), else str()
                    if hasattr(s.kind, "name"):
                        s_type = s.kind.name
                    else:
                        s_type = str(s.kind)
                elif hasattr(s, "type"):
                     s_type = str(s.type)
                
                serialized.append({
                    "id": s.id,
                    "title": s.title,
                    "url": s.url,
                    "status": s.status,
                    "type": s_type
                })
            except Exception as e:
                print(f"Error serializing source {s.id}: {e}")
                # Append minimal info
                serialized.append({
                    "id": s.id,
                    "title": s.title,
                    "type": "ERROR"
                })
        return serialized
    
    async def get_source_content(self, source_id: str) -> Dict[str, Any]:
        if not self.client or not self.current_notebook_id:
            raise Exception("not_authenticated_or_selected")
        
        fulltext = await self.client.sources.get_fulltext(self.current_notebook_id, source_id)
        return {
            "content": fulltext.content,
            "title": fulltext.title,
            "url": fulltext.url,
            "char_count": fulltext.char_count
        }

    async def generate_source_summary(self, notebook_id: str, source_id: str) -> str:
        """Generate a summary and key topics for a specific source using the AI"""
        self.set_notebook(notebook_id)
        if not self.client: raise Exception("Not authenticated")
        
        # 1. Find source title
        sources = await self.get_sources()
        target_source = next((s for s in sources if s['id'] == source_id), None)
        
        if not target_source:
            print(f"Error: Source ID {source_id} not found in {len(sources)} sources")
            raise Exception(f"Source with ID {source_id} not found")
            
        title = target_source['title']
        print(f"Found target source: {title}")
        
        # 2. Query AI
        prompt = (
            f"Please analyze the source document titled '{title}'. "
            f"Provide a concise summary (2-3 sentences) and list 3-5 key topics. "
            f"Structure your response strictly with '## Summary' followed by the summary, and '## Key Topics' followed by the list. "
            f"Focus ONLY on this source."
        )
        
        print(f"Sending prompt for summary: {prompt[:50]}...")
        summary = await self.query(prompt)
        print(f"Summary generated, length: {len(summary)}")
        return summary

    async def query(self, prompt: str):
        if not self.client:
            raise Exception("Not authenticated")
        if not self.current_notebook_id:
            raise Exception("No notebook selected")
        
        self.add_message(self.current_notebook_id, "user", prompt)
        result = await self.client.chat.ask(self.current_notebook_id, prompt)
        
        # Handle result content
        answer_text = result.answer
        
        # Handle Citations if available
        citations = []
        if hasattr(result, 'citations') and result.citations:
            for c in result.citations:
                citations.append(c.content[:200] + "...") # Snippet
        
        if citations:
            answer_text += "\n\nSources used:\n- " + "\n- ".join(citations)

        self.add_message(self.current_notebook_id, "ai", answer_text)
        return answer_text

    async def get_suggested_questions(self, notebook_id: str) -> List[str]:
        """Get AI-generated suggested questions for a notebook"""
        if not self.client:
            print("get_suggested_questions: No client, attempting to connect...")
            # Try to auto-connect if not already connected
            await self.try_auto_connect()
            
        if not self.client:
            print("get_suggested_questions: Still no client after auto-connect attempt")
            return []
            
        try:
            print(f"Calling get_description for notebook {notebook_id}")
            desc = await self.client.notebooks.get_description(notebook_id)
            print(f"Description received: summary={desc.summary[:100] if desc.summary else 'None'}...")
            print(f"Suggested topics count: {len(desc.suggested_topics)}")
            for i, topic in enumerate(desc.suggested_topics):
                print(f"  Topic {i}: question='{topic.question}', prompt='{topic.prompt[:50] if topic.prompt else 'None'}...'")
            questions = [t.question for t in desc.suggested_topics if t.question]
            print(f"Returning {len(questions)} questions")
            return questions
        except Exception as e:
            print(f"Error getting suggestions: {e}")
            import traceback
            traceback.print_exc()
            return []

    async def stream_query(self, prompt: str):
        """Stream the response from NotebookLM in real-time (simulated)"""
        if not self.client:
            raise Exception("Not authenticated")
        if not self.current_notebook_id:
            raise Exception("No notebook selected")
            
        notebook_id = self.current_notebook_id
        
        # Add user message immediately
        self.add_message(notebook_id, "user", prompt)
        
        try:
            # Use the regular ask method (it's already async)
            result = await self.client.chat.ask(notebook_id, prompt)
            answer_text = result.answer
            
            # Handle Citations if available
            citations = []
            if hasattr(result, 'citations') and result.citations:
                for i, c in enumerate(result.citations):
                     # Attempt to get source citation content
                     content = getattr(c, 'content', '') or getattr(c, 'quote', '') or "Citation"
                     label = f"[{i+1}] {content[:50].replace('\n', ' ')}..."
                     
                     # Get source ID if available
                     sid = getattr(c, 'source_id', '')
                     if sid:
                         citations.append(f"[{label}](citation://{sid})")
                     else:
                         citations.append(label)
            
            if citations:
                answer_text += "\n\n**Sources:**\n" + "\n".join([f"- {c}" for c in citations])

            # Simulate streaming by yielding chunks of text
            # This creates a typing effect on the frontend
            chunk_size = 5  # Characters per chunk
            for i in range(0, len(answer_text), chunk_size):
                chunk = answer_text[i:i + chunk_size]
                yield chunk
                # Small delay to make the typing effect visible
                await asyncio.sleep(0.01)
            
            # Save the complete message
            self.add_message(notebook_id, "ai", answer_text)
            
        except Exception as e:
            error_msg = f"Error: {str(e)}"
            print(f"Stream query error: {e}")
            yield error_msg
            self.add_message(notebook_id, "ai", error_msg)

    # Generation methods using NotebookLM Artifacts API
    async def generate_audio(self, instructions: str = "make it engaging") -> str:
        """Generate podcast audio using NotebookLM's audio generation"""
        if not self.client or not self.current_notebook_id:
            raise Exception("Not authenticated or no notebook selected")
        
        print(f"Generating audio for notebook {self.current_notebook_id}...")
        status = await self.client.artifacts.generate_audio(
            self.current_notebook_id, 
            instructions=instructions
        )
        
        print(f"Waiting for audio generation (task_id: {status.task_id})...")
        await self.client.artifacts.wait_for_completion(
            self.current_notebook_id, 
            status.task_id,
            timeout=900
        )
        
        filename = "podcast.mp3"
        print(f"Downloading audio to {filename}...")
        await self.client.artifacts.download_audio(
            self.current_notebook_id, 
            filename
        )
        
        return os.path.abspath(filename)
    
    async def export_quiz_to_docx(self, quiz_content: str, output_path: str = None) -> str:
        """Export quiz JSON content to a Word document"""
        try:
            quiz_data = json.loads(quiz_content)
        except json.JSONDecodeError:
            raise Exception("Invalid quiz content format")

        doc = Document()
        
        # Title
        title_para = doc.add_paragraph()
        run = title_para.add_run(quiz_data.get("title", "Quiz"))
        run.bold = True
        run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph() # Spacer

        # Questions
        questions = quiz_data.get("questions", [])
        for i, q in enumerate(questions):
            # Question Text
            q_para = doc.add_paragraph(f"{i+1}. {q.get('question') or q.get('text', '')}")
            q_para.runs[0].bold = True
            
            # Options
            options = q.get("answerOptions") or q.get("options") or []
            for opt in options:
                opt_text = opt.get("text") or opt.get("option", "")
                is_correct = opt.get("isCorrect", False)
                
                # Checkbox style
                prefix = "[ ] "
                # if is_correct: prefix = "[x] " # Maybe don't reveal answer in question part?
                
                o_para = doc.add_paragraph(f"{prefix}{opt_text}")
                o_para.paragraph_format.left_indent = Pt(20)

            doc.add_paragraph() # Spacer between questions

        # Answer Key Section (New Page)
        doc.add_page_break()
        key_header = doc.add_paragraph("Answer Key")
        key_header.runs[0].bold = True
        key_header.runs[0].font.size = Pt(14)
        
        for i, q in enumerate(questions):
            options = q.get("answerOptions") or q.get("options") or []
            correct_opt = next((opt for opt in options if opt.get("isCorrect")), None)
            
            if correct_opt:
                correct_text = correct_opt.get("text") or correct_opt.get("option", "")
                rationale = correct_opt.get("rationale", "")
                
                p = doc.add_paragraph()
                p.add_run(f"{i+1}. {correct_text}").bold = True
                if rationale:
                    p.add_run(f"\nRationale: {rationale}").italic = True
                    
            doc.add_paragraph()

        filename = output_path or "quiz_export.docx"
        doc.save(filename)
        return os.path.abspath(filename)

    async def generate_video(self, style: str = "whiteboard") -> str:
        """Generate video using NotebookLM's video generation"""
        if not self.client or not self.current_notebook_id:
            raise Exception("Not authenticated or no notebook selected")
        
        print(f"Generating video for notebook {self.current_notebook_id}...")
        status = await self.client.artifacts.generate_video(
            self.current_notebook_id,
            style=style
        )
        
        print(f"Waiting for video generation (task_id: {status.task_id})...")
        await self.client.artifacts.wait_for_completion(
            self.current_notebook_id,
            status.task_id,
            timeout=900
        )
        
        filename = "video.mp4"
        print(f"Downloading video to {filename}...")
        await self.client.artifacts.download_video(
            self.current_notebook_id,
            filename
        )
        
        return os.path.abspath(filename)
    
    async def generate_quiz(self, difficulty: str = "medium", quantity: str = "standard", instructions: str = None, output_format: str = "json") -> str:
        """Generate quiz using NotebookLM's quiz generation"""
        if not self.client or not self.current_notebook_id:
            raise Exception("Not authenticated or no notebook selected")
        
        print(f"Generating quiz for notebook {self.current_notebook_id}...")
        
        # Ensure Traditional Chinese is used by default if not specified
        if not instructions:
            instructions = "請使用「繁體中文」生成測驗內容。"
        elif "中文" not in instructions and "Chinese" not in instructions:
            instructions = f"請使用「繁體中文」生成：{instructions}"
            
        # Map difficulty string to enum
        difficulty_map = {
            "easy": QuizDifficulty.EASY,
            "medium": QuizDifficulty.MEDIUM,
            "hard": QuizDifficulty.HARD
        }
        diff_enum = difficulty_map.get(difficulty.lower(), QuizDifficulty.MEDIUM)
        
        # Map quantity string to enum
        quantity_map = {
            "fewer": QuizQuantity.FEWER,
            "standard": QuizQuantity.STANDARD
        }
        qty_enum = quantity_map.get(quantity.lower(), QuizQuantity.STANDARD)
        
        status = await self.client.artifacts.generate_quiz(
            self.current_notebook_id,
            difficulty=diff_enum,
            quantity=qty_enum,
            instructions=instructions
        )
        
        print(f"Waiting for quiz generation (task_id: {status.task_id})...")
        await self.client.artifacts.wait_for_completion(
            self.current_notebook_id,
            status.task_id,
            timeout=900
        )
        
        ext = "json" if output_format == "json" else "md"
        filename = f"quiz.{ext}"
        print(f"Downloading quiz to {filename}...")
        await self.client.artifacts.download_quiz(
            self.current_notebook_id,
            filename,
            output_format=output_format
        )
        
        # Post-process the file to ensure it's human-readable (not escaped Unicode)
        if output_format == "json" and os.path.exists(filename):
            try:
                content = json.loads(Path(filename).read_text(encoding="utf-8"))
                Path(filename).write_text(
                    json.dumps(content, ensure_ascii=False, indent=2),
                    encoding="utf-8"
                )
                print(f"Optimized {filename} for readability (Trad. Chinese).")
            except Exception as e:
                print(f"Warning: Failed to post-process quiz JSON: {e}")
        
        abs_path = os.path.abspath(filename)
        # Try to read title from JSON if available
        title = "Quiz"
        try:
             if output_format == "json" and os.path.exists(filename):
                 data = json.loads(Path(filename).read_text(encoding="utf-8"))
                 title = data.get("title", "Quiz")
        except:
            pass

        self.add_artifact(self.current_notebook_id, "quiz", title, {"filename": filename, "path": abs_path, "difficulty": difficulty})
        return abs_path
    
    async def generate_mindmap(self) -> str:
        """Generate mind map using NotebookLM's mind map generation"""
        if not self.client or not self.current_notebook_id:
            raise Exception("Not authenticated or no notebook selected")
        
        print(f"Generating mind map for notebook {self.current_notebook_id}...")
        # Unofficial library's generate_mind_map is synchronous and returns a dict with 'mind_map' and 'note_id'
        data = await self.client.artifacts.generate_mind_map(
            self.current_notebook_id
        )
        
        mind_map_data = data.get("mind_map")
        if not mind_map_data:
            raise Exception("Failed to generate mind map data")
            
        filename = "mindmap.json"
        print(f"Saving mind map to {filename}...")
        
        # Save the mind map JSON data
        Path(filename).write_text(
            json.dumps(mind_map_data, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        
        abs_path = os.path.abspath(filename)
        self.add_artifact(self.current_notebook_id, "mindmap", "Mind Map", {"filename": filename, "path": abs_path})
        return abs_path
    
    async def generate_slide_deck(self) -> str:
        """Generate slide deck using NotebookLM's slide generation"""
        if not self.client or not self.current_notebook_id:
            raise Exception("Not authenticated or no notebook selected")
        
        print(f"Generating slide deck for notebook {self.current_notebook_id}...")
        status = await self.client.artifacts.generate_slide_deck(
            self.current_notebook_id
        )
        
        print(f"Waiting for slide deck generation (task_id: {status.task_id})...")
        await self.client.artifacts.wait_for_completion(
            self.current_notebook_id,
            status.task_id,
            timeout=900
        )
        
        filename = "slides.pdf"
        print(f"Downloading slide deck to {filename}...")
        await self.client.artifacts.download_slide_deck(
            self.current_notebook_id,
            filename
        )
        
        abs_path = os.path.abspath(filename)
        self.add_artifact(self.current_notebook_id, "slides", "Slide Deck", {"filename": filename, "path": abs_path})
        return abs_path

    async def generate_study_guide(self) -> str:
        """Generate study guide using NotebookLM's study guide generation"""
        if not self.client or not self.current_notebook_id:
            raise Exception("Not authenticated or no notebook selected")
        
        print(f"Generating study guide for notebook {self.current_notebook_id}...")
        
        # Use generate_report with Traditional Chinese prompt and language
        custom_prompt = (
            "請為這份 Notebook 的內容產生一份完整的學習指南（Study Guide）。\n"
            "要求：\n"
            "- 必須完全使用「繁體中文」撰寫。\n"
            "- 包含術語表 (Glossary)、問答題 (Q&A) 以及重點摘要。\n"
            "- 內容必須嚴謹且準確地反映原始文件。"
        )
        
        status = await self.client.artifacts.generate_report(
            self.current_notebook_id,
            report_format=ReportFormat.STUDY_GUIDE,
            language="zh-TW",
            custom_prompt=custom_prompt
        )
        
        print(f"Waiting for study guide generation (task_id: {status.task_id})...")
        await self.client.artifacts.wait_for_completion(
            self.current_notebook_id,
            status.task_id,
            timeout=900
        )
        
        filename = "study_guide.md"
        print(f"Downloading study guide to {filename}...")
        await self.client.artifacts.download_report(
            self.current_notebook_id,
            filename,
            artifact_id=status.task_id
        )
        
        abs_path = os.path.abspath(filename)
        self.add_artifact(self.current_notebook_id, "study_guide", "Study Guide", {"filename": filename, "path": abs_path, "language": "zh-TW"})
        return abs_path
    
    async def generate_flashcards(self, quantity: str = "normal", output_format: str = "json") -> str:
        """Generate flashcards using NotebookLM's flashcard generation"""
        if not self.client or not self.current_notebook_id:
            raise Exception("Not authenticated or no notebook selected")
        
        print(f"Generating flashcards for notebook {self.current_notebook_id}...")
        
        # Map quantity string to enum
        quantity_map = {
            "less": QuizQuantity.FEWER,
            "normal": QuizQuantity.STANDARD,
            "more": QuizQuantity.STANDARD # Library only has FEWER/STANDARD
        }
        qty_enum = quantity_map.get(quantity.lower(), QuizQuantity.STANDARD)
        
        status = await self.client.artifacts.generate_flashcards(
            self.current_notebook_id,
            quantity=qty_enum
        )
        
        print(f"Waiting for flashcards generation (task_id: {status.task_id})...")
        await self.client.artifacts.wait_for_completion(
            self.current_notebook_id,
            status.task_id,
            timeout=900
        )
        
        ext = "json" if output_format == "json" else "md"
        filename = f"flashcards.{ext}"
        print(f"Downloading flashcards to {filename}...")
        await self.client.artifacts.download_flashcards(
            self.current_notebook_id,
            filename,
            output_format=output_format
        )
        
        abs_path = os.path.abspath(filename)
        self.add_artifact(self.current_notebook_id, "flashcards", "Flashcards", {"filename": filename, "path": abs_path})
        return abs_path

manager = NotebookManager()
